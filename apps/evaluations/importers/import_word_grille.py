# apps/evaluations/importers/import_word_grille.py
"""
Module pour importer des grilles d'évaluation depuis des documents Word
"""
import re
from typing import Dict, List, Optional
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.core.exceptions import ValidationError
from apps.evaluations.models_grilles import (
    GrilleEvaluation,
    CritereEvaluation,
    ElementEvaluation,
    TypeGrilleEvaluation
)
from apps.utilisateurs.models_formation import Cours, Classe


class WordGrilleImporter:
    """Classe pour importer une grille d'évaluation depuis un document Word"""
    
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.doc = Document(file_path)
        self.grille_data = {}
    
    def parse_document(self) -> Dict:
        """
        Parse le document Word et extrait la structure de la grille
        
        Format attendu dans le Word :
        - Titre de la grille (premier paragraphe ou titre 1)
        - Description (paragraphe suivant)
        - Section "CRITÈRES" ou "CRITÈRES D'ÉVALUATION"
        - Pour chaque critère :
          * Numéro/ordre et libellé
          * Description optionnelle
          * Éléments (sous-critères) avec indentation ou numérotation
        """
        grille_data = {
            'titre': '',
            'description': '',
            'criteres': []
        }
        
        current_section = None
        current_critere = None
        
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Détecter le titre (premier paragraphe non vide ou style Heading 1)
            if not grille_data['titre'] and (para.style.name.startswith('Heading') or len(text) < 100):
                grille_data['titre'] = text
                continue
            
            # Détecter la description
            if grille_data['titre'] and not grille_data['description'] and not text.upper().startswith('CRITÈRE'):
                grille_data['description'] = text
                continue
            
            # Détecter la section Critères
            if 'CRITÈRE' in text.upper() or 'ÉVALUATION' in text.upper():
                current_section = 'criteres'
                continue
            
            # Parser les critères
            if current_section == 'criteres':
                # Détecter un nouveau critère (numérotation : 1., 1), I., etc.)
                critere_match = re.match(r'^(\d+)[\.\)]\s*(.+)', text)
                if critere_match:
                    # Sauvegarder le critère précédent
                    if current_critere:
                        grille_data['criteres'].append(current_critere)
                    
                    # Nouveau critère
                    ordre = int(critere_match.group(1))
                    libelle = critere_match.group(2).strip()
                    current_critere = {
                        'ordre': ordre,
                        'libelle': libelle,
                        'description': '',
                        'poids': 1.0,
                        'elements': []
                    }
                # Détecter un élément (sous-critère) - indentation ou numérotation secondaire
                elif current_critere and (text.startswith('  ') or text.startswith('\t') or 
                                         re.match(r'^[a-z]\)|^[ivx]+\)|^-\s', text, re.IGNORECASE)):
                    element_text = re.sub(r'^[a-z]\)\s*|^[ivx]+\)\s*|^-\s*', '', text, flags=re.IGNORECASE).strip()
                    element_text = element_text.lstrip(' \t')
                    if element_text:
                        current_critere['elements'].append({
                            'libelle': element_text,
                            'description': '',
                            'poids': 1.0
                        })
                # Description du critère (paragraphe suivant sans numérotation)
                elif current_critere and not re.match(r'^\d+[\.\)]', text):
                    if not current_critere['description']:
                        current_critere['description'] = text
                    else:
                        current_critere['description'] += '\n' + text
        
        # Ajouter le dernier critère
        if current_critere:
            grille_data['criteres'].append(current_critere)
        
        self.grille_data = grille_data
        return grille_data
    
    def create_grille(self, type_grille_id: int, cours_id: Optional[int] = None, 
                     classe_id: Optional[int] = None, createur=None) -> GrilleEvaluation:
        """Crée une grille d'évaluation à partir des données parsées"""
        if not self.grille_data:
            self.parse_document()
        
        type_grille = TypeGrilleEvaluation.objects.get(pk=type_grille_id)
        
        grille = GrilleEvaluation.objects.create(
            type_grille=type_grille,
            titre=self.grille_data.get('titre', 'Grille importée'),
            description=self.grille_data.get('description', ''),
            cours_id=cours_id,
            classe_id=classe_id,
            createur=createur,
            note_maximale=20.0,
            echelle_evaluation='1-5',
            actif=True
        )
        
        # Créer les critères
        for critere_data in self.grille_data.get('criteres', []):
            critere = CritereEvaluation.objects.create(
                grille=grille,
                ordre=critere_data.get('ordre', 0),
                libelle=critere_data['libelle'],
                description=critere_data.get('description', ''),
                poids=critere_data.get('poids', 1.0),
                actif=True
            )
            
            # Créer les éléments
            for idx, element_data in enumerate(critere_data.get('elements', []), 1):
                ElementEvaluation.objects.create(
                    critere=critere,
                    ordre=idx,
                    libelle=element_data['libelle'],
                    description=element_data.get('description', ''),
                    poids=element_data.get('poids', 1.0),
                    actif=True
                )
        
        return grille


def create_word_template(output_path: str):
    """
    Crée un modèle Word pour les grilles d'évaluation
    
    Format du template :
    - Titre de la grille
    - Description
    - Section Critères avec numérotation
    - Pour chaque critère : libellé, description, éléments
    """
    doc = Document()
    
    # Titre
    title = doc.add_heading('GRILLE D\'ÉVALUATION', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Instructions
    doc.add_paragraph('Modèle pour créer une grille d\'évaluation')
    doc.add_paragraph('Remplissez les sections ci-dessous et importez le document dans le système.')
    doc.add_paragraph('')
    
    # Section Titre de la grille
    doc.add_heading('1. INFORMATIONS GÉNÉRALES', level=1)
    doc.add_paragraph('Titre de la grille : [À remplir]')
    doc.add_paragraph('Description : [À remplir]')
    doc.add_paragraph('')
    
    # Section Critères
    doc.add_heading('2. CRITÈRES D\'ÉVALUATION', level=1)
    doc.add_paragraph('Listez les critères d\'évaluation avec leur numérotation :')
    doc.add_paragraph('')
    
    # Exemple de critère
    doc.add_paragraph('1. [Libellé du critère 1]', style='List Number')
    doc.add_paragraph('   Description du critère (optionnel)')
    doc.add_paragraph('   a) Élément 1 du critère')
    doc.add_paragraph('   b) Élément 2 du critère')
    doc.add_paragraph('')
    
    doc.add_paragraph('2. [Libellé du critère 2]', style='List Number')
    doc.add_paragraph('   Description du critère (optionnel)')
    doc.add_paragraph('   a) Élément 1 du critère')
    doc.add_paragraph('   b) Élément 2 du critère')
    doc.add_paragraph('')
    
    doc.add_paragraph('3. [Libellé du critère 3]', style='List Number')
    doc.add_paragraph('   Description du critère (optionnel)')
    doc.add_paragraph('')
    
    # Instructions de formatage
    doc.add_heading('3. INSTRUCTIONS', level=1)
    instructions = [
        'Utilisez la numérotation automatique (1., 2., 3., etc.) pour les critères principaux',
        'Utilisez des lettres minuscules (a), b), c), etc.) ou des tirets (-) pour les éléments/sous-critères',
        'L\'indentation peut être faite avec des espaces ou des tabulations',
        'Chaque critère peut avoir une description optionnelle sur la ligne suivante',
        'Les éléments sont optionnels - un critère peut ne pas avoir d\'éléments',
        'Sauvegardez le document en format .docx',
        'Importez-le via l\'interface d\'administration ou la commande de gestion'
    ]
    
    for instruction in instructions:
        doc.add_paragraph(instruction, style='List Bullet')
    
    # Sauvegarder
    doc.save(output_path)
    return output_path

